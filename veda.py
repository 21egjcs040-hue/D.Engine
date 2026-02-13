import os
import sys
import subprocess
import io
import math
import base64
from PIL import Image
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

# --- HELPER FUNCTIONS ---
def resize_logo_exact(image_path, target_width_in, target_height_in):
    DPI = 96  # Word safe DPI

    target_w_px = int(target_width_in * DPI)
    target_h_px = int(target_height_in * DPI)

    img = Image.open(image_path)
    img = img.convert("RGBA")
    img = img.resize((target_w_px, target_h_px), Image.Resampling.LANCZOS)

    bio = io.BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    return bio

def process_logo(img_file, max_height_inches=0.55, max_width_inches=1.5):
    img = Image.open(img_file)
    width, height = img.size

    # Pillow 10+ compatible resampling
    try:
        resample_method = Image.Resampling.LANCZOS
    except AttributeError:
        resample_method = Image.LANCZOS  # Image.ANTIALIAS ab remove ho gaya


    # ratio calculate karo
    ratio = min(max_width_inches / (width/96), max_height_inches / (height/96), 1)
    new_width = int(width * ratio)
    new_height = int(height * ratio)

    img = img.resize((new_width, new_height), resample=resample_method)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def add_page_border(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    sec = doc.sections[0]
    sectPr = sec._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        

        border_el.set(qn('w:sz'), '4') 
        
        border_el.set(qn('w:col'), '000000')  # Black color
        border_el.set(qn('w:space'), '24')    # Gap from page edge
        pgBorders.append(border_el)
    
    sectPr.append(pgBorders)


def create_streamlit_config():
    config_dir = ".streamlit"
    if not os.path.exists(config_dir):
        os.makedirs(config_dir)
    
    config_path = os.path.join(config_dir, "config.toml")
    
    config_content = """
[theme]
primaryColor = "#D81B60"
backgroundColor = "#FFF0F5"
secondaryBackgroundColor = "#FFE4E1"
textColor = "#000000"
font = "sans serif"
    """
    with open(config_path, "w") as f:
        f.write(config_content.strip())

if __name__ == "__main__":
    create_streamlit_config()
    if os.environ.get('STREAMLIT_IS_RUNNING') != 'true':
        print("🚀 start the tool... Please wait...")
        os.environ['STREAMLIT_IS_RUNNING'] = 'true'
        try:
            subprocess.run(["streamlit", "run", sys.argv[0]], check=True)
        except Exception:
            subprocess.run([sys.executable, "-m", "streamlit", "run", sys.argv[0]])
        sys.exit()


def add_custom_footer(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Footer ka paragraph saaf karke naya banayein
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear() 
    
    # --- LEFT SIDE: Company Name ---
    run_left = p.add_run(
    "1. Should match the latest sanctioned plan\n"
    "2. As per IGBC eligibility condition for IGBC GAH"
    )
    run_left.font.size = Pt(7)
    run_left.font.bold = True
    run_left.font.color.rgb = RGBColor(0, 0, 0)

    # --- BEECH MEIN SPACE (Tab) ---
    # Ye \t text ko right side dhakelta hai
    p.add_run("\t\t\t\t\t\t") 

    # --- RIGHT SIDE: Page Number ---
    run_right = p.add_run("Page ")
    run_right.font.size = Pt(10)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_right._r.append(fldChar1)
    run_right._r.append(instrText)
    run_right._r.append(fldChar2)

    # Paragraph ki alignment LEFT hi rakhein, Tab khud right manage karega
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def make_row_cant_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), 'true')
    trPr.append(cantSplit)



def process_image_for_word(img_file,
                           land_w=4.25, land_h=2.25,
                           port_w=3.6, port_h=2.03):
    try:
        img = Image.open(img_file)
        img = img.convert("RGB")

        w_px, h_px = img.size
        aspect = w_px / h_px
        DPI = 96

        # 🎯 LANDSCAPE IMAGE
        if aspect >= 1:
            max_w_px = int(land_w * DPI)
            max_h_px = int(land_h * DPI)

            scale = min(max_w_px / w_px, max_h_px / h_px)
        
        # 🎯 PORTRAIT IMAGE
        else:
            max_w_px = int(port_w * DPI)
            max_h_px = int(port_h * DPI)

            scale = min(max_w_px / w_px, max_h_px / h_px)

        new_w = int(w_px * scale)
        new_h = int(h_px * scale)

        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf

    except Exception:
        return None



def inject_custom_logo(logo_filename):
    if os.path.exists(logo_filename):
        try:
            st.logo(logo_filename, icon_image=logo_filename)
        except:
            pass 

        with open(logo_filename, "rb") as f:
            data = f.read()
        encoded = base64.b64encode(data).decode()
        
        st.markdown(
            f"""
            <style>
                [data-testid="stHeader"] {{
                    background-color: #FFF0F5 !important;
                    z-index: 1;
                }}
                [data-testid="stHeader"]::before {{
                    content: "";
                    background-image: url("data:image/png;base64,{encoded}");
                    background-repeat: no-repeat;
                    background-size: contain;
                    background-position: center;
                    position: absolute;
                    right: 180px; 
                    top: 10px;
                    width: 60px;
                    height: 50px;
                    z-index: 999;
                    opacity: 1 !important;
                }}
            </style>
            """,
            unsafe_allow_html=True,
        )

inject_custom_logo("kamal logo.jpg") 

# 1. Page Config
st.set_page_config(page_title="D-Engine 2.0 | Kamal Cogent Energy", layout="wide", page_icon="🌿")

# --- LOGO IN TOP INTERFACE ---


# --- CSS STYLING ---
st.markdown("""
<style>
    body, .stApp, div { font-family: 'Segoe UI', sans-serif; }
    .stMarkdown, p, h1, h2, h3, h4, h5, h6, label, span, div, li { color: #000000 !important; }
    
    input, textarea, .stDateInput > div > div, div[data-baseweb="select"] > div {
        background-color: #ffffff !important;
        color: #000000 !important;
        border: 1px solid #D81B60 !important;
        border-radius: 5px !important;
    }
    div[data-baseweb="select"] div { color: #000000 !important; -webkit-text-fill-color: #000000 !important; }
    div[data-baseweb="select"] svg { fill: #000000 !important; }
    div[data-baseweb="menu"], div[data-baseweb="popover"], div[data-baseweb="option"] {
        background-color: #ffffff !important;
        color: #000000 !important;
    }
    
    .stButton > button { background-color: #D81B60 !important; color: white !important; font-weight: bold !important; border: none; }
    .stButton > button:hover { background-color: #C2185B !important; transform: scale(1.02); }
    [data-testid="stFileUploader"] section { background-color: #ffffff !important; border: 2px dashed #D81B60 !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<h1 style='text-align: center; color: #D81B60;'> KAMAL COGENT ENERGY</h1>
<h3 style='text-align: center; color: #2E7D32;'> IGBC & LEED Compliance Report Generator</h3>
""", unsafe_allow_html=True)

ALL_IMG_TYPES = ['png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff', 'tif']

# --- SIDEBAR (Is hisse ko aise update karein) ---
with st.sidebar:
    st.markdown("### 📝 Project Details")
    p_name = st.text_input("Project Name", "Rustomjee Crown")
    p_num = st.text_input("Registration Number", "27AAA")
    p_loc = st.text_input("Location", "Mumbai")
    p_precert = st.text_input("Pre-certification", "Gold")
    p_area = st.text_input("Built-up Area", "50,000 Sq. m.")
    p_units = st.text_input("Dwelling Units", "450")
    p_afford = st.text_input("Affordable Units", "50")
    p_date = st.date_input("Report Date")

    # Tower details ko sidebar ke ANDAR lane ke liye indentation (spaces) zaroori hai
    st.markdown("---")
    st.markdown("### 🏢 Tower Details")
    num_towers = st.number_input("Add Towers?", min_value=1, max_value=20, value=1)

    towers_list = []
    for i in range(int(num_towers)):
        with st.expander(f"Tower {i+1} Configuration", expanded=True):
            col_t1, col_t2 = st.columns(2)
            with col_t1:
                t_nm = st.text_input(f"Tower Name/Number", value=f"Tower {chr(65+i)}", key=f"tnm_{i}")
                t_fl = st.text_input(f"Number of floors", value="G + 45", key=f"tfl_{i}")
            with col_t2:
                t_st = st.text_input(f"Construction stage (%)", value="40%", key=f"tst_{i}")
            towers_list.append({"name": t_nm, "floors": t_fl, "stage": t_st})

    # Header Titles wala part bhi sidebar ke andar hi rahega (Indented)
    st.markdown("---")
    st.markdown("### 🖼️ Header & Titles")
    col_h1, col_h2 = st.columns(2)
    with col_h1: logo_left = st.file_uploader("Left Logo", type=ALL_IMG_TYPES)
    with col_h2: logo_right = st.file_uploader("Right Logo", type=ALL_IMG_TYPES)
    header_center_text = st.text_area("Header Center Text", "", height=70)
    main_body_title = st.text_input("Main Report Title", "IGBC Six Monthly Compliance Report")
    
    header_color_input = st.color_picker("Pick Color", "#48B448")
    uploaded_template = st.file_uploader("Upload .docx Template", type=['docx'])

# Caption Options
caption_options = [
    "Existing site conditions", "Topsoil preservation, etc", "Green cover on site & irrigation system", 
    "Heat Island effect (Roof)", "Heat Island effect (non-roof)", "Differently abled facilities", 
    "Basic facilities for construction workforce", "Electric charging facility", "STP installation & dual plumbing system", 
    "Rainwater Harvesting / recharge pit construction", 
    "Energy efficient building envelope  •	Wall construction  •	Roof construction   •	Glass    •	Projections for openings", 
    "Renewable energy & Hot water system", 
    "Segregation of construction waste, reuse applications, gatepass/ challans of materials sold", "Use of local materials (photos & Invoices/ delivery challans)",
    "Material with recycled content (photos & Invoices/ delivery challans)", 
    "Paints & adhesives (photos & Invoices/ delivery challans)", "Alternate construction materials (photos & Invoices/ delivery challans)",
    "Organic Waste convertor and space requirements","Any other relevant information","Reference Photographs"
]
status_options = ["Completed", "In Progress", "To be initiated", "Pending"]

# --- MAIN UPLOAD ---
st.info("💡 **Feature:** All image formats supported. Auto-groups same captions.")

uploaded_files = st.file_uploader("📂 Upload MAIN Photos", type=ALL_IMG_TYPES, accept_multiple_files=True)

entries_data = {}

if uploaded_files:
    st.write("---")
    for i, file1 in enumerate(uploaded_files):
        with st.container():
            c_check, c_img, c_ctrl = st.columns([0.5, 2.5, 5])
            with c_check:
                st.write("")
                st.write("")
                include = st.checkbox(f"#{i+1}", value=True, key=f"chk_{i}")
            
            if include:
                with c_img:
                    st.image(file1, use_column_width=True, caption=f"Main Photo {i+1}")
                with c_ctrl:
                    cc1, cc2 = st.columns(2)
                    with cc1:
                        select_options = ["Select Caption...", "➕ Add Custom Caption..."] + caption_options
                        selected_option = st.selectbox(f"Caption", select_options, key=f"c_sel_{i}")
                        if selected_option == "➕ Add Custom Caption...":
                            final_caption = st.text_input("Apna Caption Likhein:", key=f"c_input_{i}")
                        else:
                            final_caption = selected_option
                    with cc2:
                        stat = st.selectbox(f"Status", status_options, key=f"s_{i}")
                    
                    st.markdown("**Add Grid Photos:**")
                    sec_files = st.file_uploader(f"Side images #{i+1}", type=ALL_IMG_TYPES, key=f"sec_{i}", accept_multiple_files=True)
                    if sec_files:
                        p_cols = st.columns(min(len(sec_files), 4) if sec_files else 1)
                        for idx, sf in enumerate(sec_files):
                            with p_cols[idx % len(p_cols)]:
                                st.image(sf, width=80)

                    if final_caption and final_caption != "Select Caption...":
                        entries_data[i] = {
                            "caption": final_caption, "status": stat, "img1": file1, "sec_imgs": sec_files
                        }
            else:
                with c_img: st.caption("Skipped")
        st.markdown("---")

    if st.button("✅ GENERATE REPORT"):
        if not entries_data:
            st.error("⚠️ Please select captions.")
        else:
            with st.spinner("Compiling Report..."):
                try:
                    if uploaded_template:
                        doc = Document(uploaded_template)
                        doc.add_paragraph("")
                    else:
                        doc = Document()
                        section = doc.sections[0]
                        section.left_margin = Inches(0.5)
                        section.right_margin = Inches(0.5)
                        section.top_margin = Inches(0.5) 
                        # ✅ Header ko page ke bilkul upar lao
                        section.header_distance = Inches(0.3)


                        # --- HEADER LOGIC ---

                        header = section.header
                        header.is_linked_to_previous = False
                        for paragraph in header.paragraphs:
                            p = paragraph._element
                            p.getparent().remove(p)

                        # Table ki total width 7.5 inches rakhein (A4 standard with 0.5 margins)
                        htable = header.add_table(rows=1, cols=3, width=Inches(7.5))

                        htable.allow_autofit = False # Ye line logo ko bahar jaane se rokegi

                        # Columns ki width ko strictly fix karein
                        htable.columns[0].width = Inches(1.5) # Left Logo space
                        htable.columns[1].width = Inches(4.5) # Center Text space
                        htable.columns[2].width = Inches(1.5) # Right Logo space

                        from docx.enum.table import WD_ROW_HEIGHT_RULE

                        htable.rows[0].height = Inches(0.8)  # Row height fix
                        htable.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


                        def set_cell_padding(cell, top=100, start=100, bottom=100, end=100):
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcMar = OxmlElement('w:tcMar')
                            for side, val in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
                                node = OxmlElement(f'w:{side}')
                                node.set(qn('w:w'), str(val))
                                node.set(qn('w:type'), 'dxa')
                                tcMar.append(node)
                            tcPr.append(tcMar)

                        # Sabhi cells mein thoda padding add karein (approx 0.05-0.1 inch)
                        for i in range(3):
                            set_cell_padding(htable.cell(0, i), top=150, bottom=150, start=100, end=100)
                            htable.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            # Paragraph ki extra spacing remove karein
                            p = htable.cell(0, i).paragraphs[0]
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = 1.0

                        # 1. Left Logo (Height constrain karein)
                        if logo_left:
                            cell = htable.cell(0, 0)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run()
                            p.paragraph_format.space_before = Pt(6)
                            p.paragraph_format.space_after = Pt(6)
                            img = resize_logo_exact(logo_left, 0.5, 0.5)
                            run.add_picture(img)


                            

                        # 2. Center Text
                        cell = htable.cell(0, 1)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if header_center_text:
                            run = p.add_run(header_center_text)
                            run.bold = True
                            run.font.size = Pt(9)

                        # 3. Right Logo (Height constrain karein)
                        if logo_right:
                            cell = htable.cell(0, 2)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            run = p.add_run()
                            p.paragraph_format.space_before = Pt(6)
                            p.paragraph_format.space_after = Pt(6)

                            img = resize_logo_exact(logo_right, 0.5, 0.5)
                            run.add_picture(img)


                            

                            # Header ke baad minimal gap
                            p_gap = doc.add_paragraph("")
                            p_gap.paragraph_format.space_before = Pt(0)
                            p_gap.paragraph_format.space_after = Pt(2)

                            # ✅ REMOVE EXTRA SPACE AFTER HEADER
                    final_header_color = header_color_input.replace('#', '')

                    # --- MAIN REPORT TITLE (BODY) ---
                    if main_body_title:
                        try:
                            # Title ki jagah Heading 1 (level=1) use kar rahe hain jo ki safe hai
                            head = doc.add_heading("", level=1)
                            run = head.add_run(main_body_title)
                            head.paragraph_format.space_before = Pt(0)
                            head.paragraph_format.space_after = Pt(4)
                            run.font.size = Pt(18)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.bold = True
                        except Exception:
                            # Agar Heading style fail ho jaye toh manual paragraph format
                            head = doc.add_paragraph()
                            run = head.add_run(main_body_title)
                            run.bold = True
                            run.font.size = Pt(32)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        # Heading ko center mein align karein
                        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        
                    
                    # --- PROJECT INFO TABLE ---
                    doc.add_paragraph("")
                    info_table = doc.add_table(rows=0, cols=2)
                    info_table.style = 'Table Grid'
                    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    details = [
                        ("Project Name", p_name),
                        ("Registration Number", p_num),   # ✅ YE LINE ADD KARO
                        ("Location", p_loc),
                        ("Pre-certification achieved", p_precert),
                        ("Total built-up area", p_area),
                        ("Total number of Dwelling Units", p_units),
                        ("Number of unit (Affordable)", p_afford),
                        ("Date", str(p_date))
                    ]
                    for lab, val in details:
                        row = info_table.add_row()
                        row.cells[0].width = Inches(2.0)
                        row.cells[1].width = Inches(4.0)
                        set_cell_background(row.cells[0], final_header_color)
                        row.cells[0].text = lab
                        p = row.cells[0].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.runs[0]
                        run.font.bold = True
                        run.font.size = Pt(9)
                        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,0)
                        row.cells[1].text = val
                        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                   # --- TOWER INFO TABLE (IS POORE BLOCK KO REPLACE KAREIN) ---
                    doc.add_paragraph("")
                    tower_table = doc.add_table(rows=2, cols=3)
                    tower_table.style = 'Table Grid'
                    tower_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # 1. Main Merged Header (Construction status as on date)
                    main_header_cell = tower_table.rows[0].cells[0].merge(tower_table.rows[0].cells[2])
                    main_header_cell.text = f"Construction status as on ({p_date.strftime('%d-%m-%Y')})"
                    set_cell_background(main_header_cell, final_header_color)

                    # --- GAP KAM KARNE KE LIYE FORMATTING ---
                    p = main_header_cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0) # Table ke andar upar ka gap khatam
                    p.paragraph_format.space_after = Pt(0)  # Table ke andar niche ka gap khatam

                    # Text ko bold karne ka sahi tarika
                    if p.runs:
                        p.runs[0].font.bold = True

                    # 2. Sub-headers Labels
                    sub_headers = [
                        "Tower name/number\n(For eg: Tower A, Tower A1 etc.)", 
                        "Number of floors\n(For eg.: G+1, S +3 etc.)", 
                        "Construction stage (%)"
                    ]

                    for i, h_text in enumerate(sub_headers):
                        cell = tower_table.rows[1].cells[i]
                        cell.text = h_text
                        set_cell_background(cell, final_header_color)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p.runs:
                            p.runs[0].font.bold = True
                            p.runs[0].font.size = Pt(10)

                    # 3. Data Rows Loop (Yahan 'vals' wali error solve hogi)
                    for t in towers_list:
                        row = tower_table.add_row()
                        row.cells[0].text = t["name"]
                        row.cells[1].text = t["floors"]
                        row.cells[2].text = t["stage"]
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    doc.add_paragraph("")

                    # --- MAIN TABLE (AUTO-GROUPING) ---
                    grouped_entries = {}
                    for entry in entries_data.values():
                        key = (entry['caption'], entry['status'])
                        if key not in grouped_entries: 
                            grouped_entries[key] = []
                        grouped_entries[key].append(entry['img1'])
                        if entry['sec_imgs']: 
                            grouped_entries[key].extend(entry['sec_imgs'])

                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    h_cells = table.rows[0].cells
                    h_cells[0].width = Inches(1.2)
                    h_cells[1].width = Inches(0.8)
                    h_cells[2].width = Inches(4.5)
                    
                    col_headers = ['Credit', 'Implementation Status (For e.g.: to be initiated,in progress, Completed)', 'Time stamp Photograph with caption']
                    for i, txt in enumerate(col_headers):
                        h_cells[i].text = txt
                        set_cell_background(h_cells[i], final_header_color)
                        p = h_cells[i].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.runs[0].font.bold = True
                        p.runs[0].font.color.rgb = RGBColor(0,0,0)
                    
                    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))

                    
                        
# --- 1. Fixed Captions (Standard List) ---
                    # हम पहले 'caption_options' के हिसाब से फोटो लगाएंगे ताकि क्रम (Order) सही रहे
# --- 1. Fixed Captions (Sync Logic) ---
                    for fixed_cap in caption_options:
                        relevant_keys = [k for k in grouped_entries.keys() if k[0] == fixed_cap]



                            # ⭐ NEW CODE ADD KAREIN (EMPTY CAPTION ROW)
                        if not relevant_keys:
                            row = table.add_row()
                            row.cells[0].text = fixed_cap
                            row.cells[1].text = ""   # status blank
                            row.cells[2].text = ""   # photo blank

                            for c in row.cells:
                                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                p = c.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                if p.runs:
                                    p.runs[0].font.size = Pt(9)
                                    p.runs[0].font.bold = True
                            continue



                        
                        for key in relevant_keys:
                            cap_text, stat_text = key
                            image_list = grouped_entries[key]
                            
                            # फोटो के 2-2 के जोड़े बनाना ताकि कैप्शन साथ रहे
                            chunk_size = 2
                            chunks = [image_list[i:i + chunk_size] for i in range(0, len(image_list), chunk_size)]
                            
                            cells_to_merge_cap = []
                            cells_to_merge_stat = []

                            for idx, chunk in enumerate(chunks):
                                row = table.add_row()
                                # रो को फटने से रोकना
                                row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
                                
                                cells_to_merge_cap.append(row.cells[0])
                                cells_to_merge_stat.append(row.cells[1])
                                
                                photo_cell = row.cells[2]
                                photo_cell.paragraphs[0].clear()
                                
                                # अंदर की ग्रिड टेबल
                                inner_table = photo_cell.add_table(rows=1, cols=len(chunk))
                                inner_table.allow_autofit = False
                                inner_table.width = Inches(4.5)

                                for c_idx, img_data in enumerate(chunk):
                                    cell = inner_table.cell(0, c_idx)
                                    cell_p = cell.paragraphs[0]
                                    cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    img_stream = process_image_for_word(img_data, 2.1, 1.8)
                                    if img_stream:
                                        run = cell_p.add_run()
                                        run.add_picture(img_stream, width=Inches(2.1))

                            # कैप्शन और स्टेटस को मर्ज करना (ASLI SOLUTION)
                            cells_to_merge_cap[0].text = str(cap_text)
                            cells_to_merge_stat[0].text = str(stat_text)

                            # अलाइनमेंट ठीक करने के लिए (CENTER):
                            for master_cell in [cells_to_merge_cap[0], cells_to_merge_stat[0]]:
                                master_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # ऊपर-नीचे से बीच में
                                p = master_cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # दाएं-बाएं से बीच में
                                if p.runs:
                                    p.runs[0].font.size = Pt(9)
                                    p.runs[0].font.bold = True

                            if len(cells_to_merge_cap) > 1:
                                for i in range(1, len(cells_to_merge_cap)):
                                    cells_to_merge_cap[0].merge(cells_to_merge_cap[i])
                                    cells_to_merge_stat[0].merge(cells_to_merge_stat[i])

                            

                            # # --- YEH HISSA ADD KAREIN ---
                            # # Isse photo wale cells ki horizontal lines hat jayengi
                            # for i in range(len(cells_to_merge_cap)):
                            #     # Photo cell (index 2) ko pakdein
                            #     photo_tc = cells_to_merge_cap[i]._tr.get_children()[2] 
                            #     tcPr = photo_tc.get_or_add_tcPr()
                            #     tcBorders = OxmlElement('w:tcBorders')
                                
                            #     # Bottom aur Top border ko 'nil' (invisible) karein
                            #     for border in ['top', 'bottom']:
                            #         node = OxmlElement(f'w:{border}')
                            #         node.set(qn('w:val'), 'nil')
                            #         tcBorders.append(node)
                            #     tcPr.append(tcBorders)

                            

                            # स्टाइलिंग: टॉप एलाइनमेंट ताकि फोटो नीचे जाने पर भी नाम ऊपर रहे
                            for master_cell in [cells_to_merge_cap[0], cells_to_merge_stat[0]]:
                                master_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                p = master_cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                if p.runs:
                                    p.runs[0].font.size = Pt(9)
                                    p.runs[0].font.bold = True

                    # --- 2. Custom Captions (वही सेम लॉजिक यहाँ भी) ---
                    custom_keys = [k for k in grouped_entries.keys() if k[0] not in caption_options]
                    for key in custom_keys:
                        cap_text, stat_text = key
                        image_list = grouped_entries[key]
                        
                        chunk_size = 2
                        chunks = [image_list[i:i + chunk_size] for i in range(0, len(image_list), chunk_size)]
                        
                        cells_to_merge_cap = []
                        cells_to_merge_stat = []

                        for idx, chunk in enumerate(chunks):
                            row = table.add_row()
                            row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
                            cells_to_merge_cap.append(row.cells[0])
                            cells_to_merge_stat.append(row.cells[1])
                            
                            photo_cell = row.cells[2]
                            photo_cell.paragraphs[0].clear()
                            inner_table = photo_cell.add_table(rows=1, cols=len(chunk))
                            inner_table.allow_autofit = False
                            inner_table.width = Inches(4.5)

                            for c_idx, img_data in enumerate(chunk):
                                cell = inner_table.cell(0, c_idx)
                                cell_p = cell.paragraphs[0]
                                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                img_stream = process_image_for_word(img_data, 2.1, 1.8)
                                if img_stream:
                                    run = cell_p.add_run()
                                    run.add_picture(img_stream, width=Inches(2.1))

                        cells_to_merge_cap[0].text = str(cap_text)
                        cells_to_merge_stat[0].text = str(stat_text)
                        if len(cells_to_merge_cap) > 1:
                            for i in range(1, len(cells_to_merge_cap)):
                                cells_to_merge_cap[0].merge(cells_to_merge_cap[i])
                                cells_to_merge_stat[0].merge(cells_to_merge_stat[i])

                        for master_cell in [cells_to_merge_cap[0], cells_to_merge_stat[0]]:
                            master_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            p = master_cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if p.runs:
                                p.runs[0].font.size = Pt(9)
                                p.runs[0].font.bold = True

                    # --- FINAL GENERATION ---
                    add_page_border(doc)
                    add_custom_footer(doc)

                    target_stream = io.BytesIO()
                    doc.save(target_stream)
                    target_stream.seek(0)

                    
                    add_custom_footer(doc)                
                    add_page_border(doc)
                    bio = io.BytesIO()
                    doc.save(bio)
                    # st.balloons()
                    st.success("🎉 Report Generated!")
                    st.download_button("📥 DOWNLOAD REPORT", bio.getvalue(), f"IGBC_{p_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                except Exception as e:
                    st.error(f"Error: {e}")
